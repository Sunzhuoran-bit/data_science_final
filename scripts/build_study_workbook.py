from __future__ import annotations

import json
import math
import textwrap
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data" / "derived"
OUT_DIR = ROOT / "docs" / "study_guides"
ASSET_DIR = OUT_DIR / "assets"

TOPIC_GUIDES = {
    "linear_diagnostics": {
        "title": "线性回归诊断与误差假设",
        "targets": ["能区分相关显著与因果判断。", "能用残差图、Q-Q 图、Cook 距离、VIF 判断模型问题。", "能说明异方差、非正态、遗漏变量和多重共线性的处理边界。"],
        "network": ["显著性", "因果判断", "残差图", "异方差", "Q-Q 图", "正态性", "Cook 距离", "强影响点", "VIF", "多重共线性"],
        "knowledge": [
            "统计显著只说明模型条件下的相关关系显著，不能直接推出因果原因。",
            "Residuals vs Fitted 图若呈喇叭状，通常指向异方差；若呈系统曲线，常提示遗漏变量或非线性结构。",
            "Q-Q 图用于观察残差是否偏离正态分布，不能把 Q-Q 偏离直接解释为异方差。",
            "Cook 距离用于识别强影响点；强影响点不等于必须删除的异常值。",
            "VIF 用于多重共线性诊断，常见经验阈值是 5 或 10，但是否剔除变量要结合模型目的。"
        ],
        "reminders": ["对数变换是常见修正方案，但不是唯一方案。", "遗漏变量偏误要同时满足遗漏变量影响 Y 且与已纳入解释变量相关。"],
    },
    "model_selection": {
        "title": "模型评价、AIC/BIC 与显著性检验",
        "targets": ["能区分 R²、调整 R²、AIC/BIC 的使用场景。", "能区分 F 检验和 t 检验。", "能完成约束 F 检验与自由度计算。"],
        "network": ["R²", "调整 R²", "AIC", "BIC", "复杂度惩罚", "F 检验", "t 检验", "自由度", "模型选择"],
        "knowledge": [
            "普通 R² 随自变量增加通常不下降，因此不能单独作为不同变量数模型的选择依据。",
            "调整 R² 引入样本量和自变量个数惩罚，适合比较同一因变量下复杂度不同的模型。",
            "AIC=-2ln(L)+2df，BIC=-2ln(L)+df ln(n)，两者均越小越好；样本量较大时 BIC 惩罚通常更强。",
            "F 检验常用于整体显著性或联合约束检验；t 检验用于单个回归系数显著性。",
            "约束 F 检验按无约束模型与约束模型的拟合差异以及自由度差进行计算。"
        ],
        "reminders": ["BIC 惩罚更强，通常保留更少变量。", "R² 大不等于模型有因果解释力或外推预测力。"],
    },
    "log_linear_house": {
        "title": "对数线性模型、交互项与二手房案例",
        "targets": ["能解释 log(y) 模型系数的百分比含义。", "能解释虚拟变量、基准组和交互项。", "能把二手房案例中的 subway、school、CATE 等变量解释清楚。"],
        "network": ["log(y)", "百分比解释", "虚拟变量", "基准组", "交互项", "城区", "学区", "单位面积房价"],
        "knowledge": [
            "log(y)=β0+β1x+ε 中，连续变量系数通常近似解释为 x 增加 1 单位时 y 的百分比变化。",
            "离散变量有 k 个水平时通常选择一个基准组，构造 k-1 个哑变量。",
            "含交互项时，主效应只是在基准条件下的效应，不能直接代表所有组的整体效应。",
            "二手房案例中，对单位面积房价取对数是为改善异方差与残差分布问题。",
            "school 系数或交互项解释必须结合控制变量和基准城区。"
        ],
        "reminders": ["不能把 log 模型系数解释成原单位的绝对金额变化。", "主效应为负不等于变量整体没有意义。"],
    },
    "anova": {
        "title": "方差分析与方差分解",
        "targets": ["能写出单因素与双因素方差分解。", "能解释 F 统计量、组间变异和组内变异。", "能说明 ANOVA 显著后的限制。"],
        "network": ["ANOVA", "SST", "SSA/SSR", "SSE", "MS", "F 统计量", "组间", "组内", "事后比较"],
        "knowledge": [
            "方差分析可视为自变量为分类变量、因变量为连续数值变量的线性模型特例。",
            "单因素 ANOVA 的核心是将总变异分解为组间变异与组内误差变异。",
            "F 值较大表示组间均方相对组内均方更大，通常支持拒绝各组均值相等的原假设。",
            "F 检验显著只能说明至少有组均值不同，不能直接判断所有组两两都不同。",
            "ANOVA 需要关注正态性、方差齐性、独立性等前提。"
        ],
        "reminders": ["显著后仍需事后比较才能判断具体哪两组不同。", "不能说只要类别自变量和数值因变量就无需假设。"],
    },
    "logistic_roc": {
        "title": "逻辑回归、ROC/AUC 与混淆矩阵",
        "targets": ["能区分 Logistic Function 与 Logit Function。", "能计算 odds、概率、TPR、FPR。", "能解释 ROC/AUC 与阈值的关系。"],
        "network": ["Logistic", "Logit", "odds", "阈值", "混淆矩阵", "TPR", "FPR", "ROC", "AUC"],
        "knowledge": [
            "逻辑回归通过 logit 把概率映射为线性预测器，系数解释为对数发生比变化。",
            "分类阈值会影响具体分类结果，但 AUC 衡量的是所有阈值下的整体排序能力。",
            "TPR=TP/(TP+FN)，FPR=FP/(FP+TN)，ROC 曲线以 FPR 为横轴、TPR 为纵轴。",
            "AUC 越接近 1，分类排序效果越好；AUC=0.5 近似随机。",
            "线性回归不适合直接处理二分类因变量，通常应使用逻辑回归等模型。"
        ],
        "reminders": ["系数 0.5 不是概率增加 50%，而是对数发生比增加 0.5。", "ROC 不是固定阈值的单点。"],
    },
    "survival_ordinal_count": {
        "title": "生存分析、定序回归与计数模型",
        "targets": ["能解释删失、Kaplan-Meier、Log-rank/Wilcoxon。", "能区分 Cox 与 AFT 的建模对象和系数方向。", "能说明定序回归阈值和泊松计数模型。"],
        "network": ["删失", "Kaplan-Meier", "Log-rank", "Wilcoxon", "Cox", "AFT", "风险比", "定序阈值", "泊松", "过离散"],
        "knowledge": [
            "生存分析处理的是事件发生时间，无法观测到确切失效时间时称为删失。",
            "Cox 模型直接建风险函数，是半参数模型；AFT 模型直接对生存时间建模，通常需指定分布。",
            "同一变量在 Cox 与 AFT 中的系数符号方向不必相同，风险增加常对应生存时间缩短。",
            "K 类定序变量通常需要 K-1 个阈值把潜变量区间切开。",
            "计数型因变量可用泊松模型，若存在过离散需考虑替代模型。"
        ],
        "reminders": ["不要把 Cox 系数直接解释成生存时间增加。", "定序编码数字不能随意当普通连续变量加减乘除。"],
    },
    "pca_fa": {
        "title": "主成分分析与因子分析",
        "targets": ["能说明 PCA 的最大方差目标。", "能解释特征值、贡献率、累计贡献率和载荷。", "能区分 PCA 与因子分析及旋转。"],
        "network": ["标准化", "协方差/相关矩阵", "特征值", "主成分", "贡献率", "载荷", "公共因子", "特殊因子", "Varimax"],
        "knowledge": [
            "PCA 寻找原始变量的线性组合，使新变量方差最大且主成分之间互不相关。",
            "贡献率按特征值除以总方差计算，累计贡献率用于判断保留多少主成分。",
            "变量量纲差异大时通常需要标准化，否则方差大的变量会主导主成分。",
            "因子分析假设观测变量由公共因子和特殊因子生成，重点解释变量间相关性的潜在来源。",
            "因子旋转用于让载荷结构更清晰；正交 Varimax 保持公共因子互不相关，通常不改变共同度。"
        ],
        "reminders": ["累计贡献率高不等于完整保留全部信息。", "公共因子通常不可直接观测。", "PCA 不负责证明变量间因果关系。"],
    },
    "clustering": {
        "title": "K-means、层次聚类与 GMM 联系",
        "targets": ["能区分聚类与分类。", "能描述 K-means 迭代步骤和局限。", "能解释层次聚类、Ward 法、树状图与 K 值选择。"],
        "network": ["无监督", "标准化", "距离", "K-means", "初始中心", "WGSS", "层次聚类", "Ward", "树状图", "业务解释"],
        "knowledge": [
            "聚类没有事先给定因变量或类别标签，属于无监督学习；分类通常属于监督学习。",
            "K-means 需要预先指定 K，反复执行分配到最近中心点与更新质心。",
            "K-means 对初始中心、异常值、变量尺度和非凸簇形状敏感，可能陷入局部最优。",
            "WGSS 碎石图可辅助选择 K；nstart 可用多组初始中心降低偶然性。",
            "层次聚类可用树状图展示合并过程，Ward 法关注类内离差平方和增量。"
        ],
        "reminders": ["K-means 不能自动推断最佳 K。", "层次聚类合并后通常不可逆，K-means 迭代中样本可重新分配。"],
    },
    "naive_bayes": {
        "title": "朴素贝叶斯与文本分类",
        "targets": ["能解释条件独立假设。", "能说明文本词频矩阵为何高维稀疏。", "能处理零概率与拉普拉斯修正。"],
        "network": ["生成式模型", "先验概率", "似然", "后验概率", "条件独立", "文本分词", "稀疏矩阵", "零概率", "拉普拉斯", "TAN"],
        "knowledge": [
            "朴素贝叶斯是生成式分类模型，核心是基于先验概率和条件似然计算后验概率。",
            "“朴素”指给定类别后特征之间条件独立，而不是现实中所有特征无相关。",
            "文本分类中，分词后可形成样本×词语矩阵，多数元素为 0，属于高维稀疏数据。",
            "训练集中某特征在某类别下从未出现会导致零概率，可用拉普拉斯修正。",
            "TAN/AODE 等半朴素方法放松完全独立假设，在复杂度和依赖表达之间折中。"
        ],
        "reminders": ["不要把朴素贝叶斯说成判别式模型。", "零概率不是只能放弃该类别，平滑可以修正。"],
    },
    "tree_ensemble": {
        "title": "决策树、随机森林与 XGBoost",
        "targets": ["能区分 ID3、C4.5、CART 的划分准则。", "能解释预剪枝、后剪枝、过拟合。", "能比较 Bagging/随机森林与 Boosting/XGBoost。"],
        "network": ["ID3", "信息增益", "C4.5", "增益率", "CART", "基尼指数", "预剪枝", "后剪枝", "随机森林", "OOB", "XGBoost", "正则化"],
        "knowledge": [
            "ID3 按信息增益选特征，C4.5 用增益率，CART 分类树常用基尼指数。",
            "树越深通常训练拟合更好，但测试集未必更好，需通过剪枝控制过拟合。",
            "预剪枝在生长过程中提前停止分支，后剪枝先生成完整树再回头修剪。",
            "随机森林通过 bootstrap 抽样和随机特征选择降低树之间相关性，可用 OOB 做无验证集评估。",
            "Bagging 通常并行训练基学习器，Boosting 串行迭代修正残差；XGBoost 在目标中加入正则项。"
        ],
        "reminders": ["Bootstrap 是有放回抽样，极限 OOB 占比约 36.8%。", "随机森林主要降低方差，Boosting 更偏向逐步降低偏差。"],
    },
    "uncategorized": {
        "title": "未归类题号与待核对内容",
        "targets": ["保留所有暂不能稳定归类的 PDF 题号。", "区分文本缺失、答案待核对和知识点待人工确认。"],
        "network": ["PDF 题号", "文本缺失", "待核对答案", "待人工归类"],
        "knowledge": [
            "本讲不新增知识，只收纳自动分类无法稳定判断或 PDF 文本提取为空的题号。",
            "这些题号保留在网页和覆盖报告中，后续可根据原 PDF 页面对题干和答案做人工核对。",
            "未归类不是丢弃；它是防止误归类、误造题的保护板块。"
        ],
        "reminders": ["不得凭空补写 PDF 文本缺失的题干。", "不得把答案待核对题用于预测卷自动评分。"],
    },
}


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_label_box(doc: Document, label: str, body: list[str], fill: str = "FFF3D8") -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.0)
    table.columns[1].width = Inches(5.3)
    label_cell, body_cell = table.rows[0].cells
    shade_cell(label_cell, "F3A43B")
    shade_cell(body_cell, fill)
    set_cell_text(label_cell, label, bold=True, color="FFFFFF", size=10)
    body_cell.text = ""
    for index, item in enumerate(body):
        p = body_cell.paragraphs[0] if index == 0 else body_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_run_font(r, size=9.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 11.5, bold=True, color="1F4D78")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_run_font(r, size=10)


def format_source(question: dict) -> str:
    source = question["source"]
    return f"{source['sourceId']} · p.{source.get('page', '?')} · {source['section']}第{source['question']}题"


def format_answer(question: dict) -> str:
    if question.get("answerPending"):
        return "待核对"
    answer = question.get("answer")
    if question["type"] == "true_false":
        return "正确" if answer else "错误"
    if isinstance(answer, list):
        return "、".join(answer)
    if isinstance(answer, dict):
        return answer.get("accepted", [""])[0]
    return str(answer)


def question_text(question: dict) -> str:
    text = question["stem"]
    options = question.get("options") or []
    if options:
        option_text = " ".join(f"{option['key']}. {option['text']}" for option in options)
        text = f"{text} {option_text}"
    return text


def make_network_image(topic_id: str, guide: dict) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1300, 420
    image = Image.new("RGB", (width, height), "#F6F8F7")
    draw = ImageDraw.Draw(image)
    font_path = next((Path(p) for p in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ] if Path(p).exists()), None)
    title_font = ImageFont.truetype(str(font_path), 28) if font_path else ImageFont.load_default()
    box_font = ImageFont.truetype(str(font_path), 22) if font_path else ImageFont.load_default()
    draw.text((40, 28), guide["title"], fill="#1F4D78", font=title_font)
    nodes = guide["network"][:10]
    cols = 5
    box_w, box_h = 210, 62
    start_x, start_y = 50, 105
    gap_x, gap_y = 38, 82
    points = []
    for index, node in enumerate(nodes):
        row = index // cols
        col = index % cols
        x = start_x + col * (box_w + gap_x)
        y = start_y + row * (box_h + gap_y)
        points.append((x + box_w / 2, y + box_h / 2))
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=14, fill="#FFFFFF", outline="#7DB59B", width=3)
        wrapped = textwrap.wrap(node, width=10)
        text_y = y + 12 if len(wrapped) == 1 else y + 5
        for line in wrapped[:2]:
            bbox = draw.textbbox((0, 0), line, font=box_font)
            draw.text((x + (box_w - (bbox[2] - bbox[0])) / 2, text_y), line, fill="#1D2520", font=box_font)
            text_y += 27
    for a, b in zip(points, points[1:]):
        draw.line((a[0], a[1], b[0], b[1]), fill="#9BAAA2", width=2)
    path = ASSET_DIR / f"{topic_id}_network.png"
    image.save(path)
    return path


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    for name, size, color in [("Heading 1", 16, "1F4D78"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 11.5, "1F4D78")]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_cover(doc: Document, questions: list[dict], topics: list[dict]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Data Science Final Review")
    set_run_font(r, size=24, bold=True, color="1F4D78")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("分知识点讲义式学案（学生版）")
    set_run_font(r, size=16, bold=True, color="2E74B5")
    add_label_box(
        doc,
        "目标要求",
        [
            "按知识点而不是 lecture 顺序组织。",
            "每个板块包含知识梳理、特别提醒、例题和 PDF 原题训练索引。",
            "所有题目均来自 4 个作业反馈 PDF；无法稳定归类的题号进入“未归类题号”。",
        ],
    )
    by_source = Counter(question["source"]["sourceId"] for question in questions)
    add_label_box(
        doc,
        "覆盖审计",
        [
            f"4 个 PDF 共保留 {len(questions)} 个唯一题号。",
            "来源分布：" + "；".join(f"{source}={count}" for source, count in sorted(by_source.items())),
            "答案待核对题仍保留题号和来源，不进入预测卷自动评分。",
        ],
        fill="EAF4EE",
    )
    add_heading(doc, "目录", 1)
    for index, topic in enumerate(topics, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"第{index}讲  {topic['name']}")
        set_run_font(run, size=10.5, bold=True)
    doc.add_page_break()


def add_examples(doc: Document, questions: list[dict]) -> None:
    verified = [question for question in questions if not question.get("answerPending")]
    examples = verified[:2] if len(verified) >= 2 else verified
    if not examples:
        add_label_box(doc, "例题", ["本板块当前只有待核对答案题，先保留原题索引，待核对后补入例题解析。"], fill="FFF3D8")
        return
    add_heading(doc, "提升 · 必考方向归纳", 2)
    for index, question in enumerate(examples, start=1):
        table = doc.add_table(rows=3, cols=1)
        table.autofit = False
        shade_cell(table.cell(0, 0), "E8EEF5")
        set_cell_text(table.cell(0, 0), f"例{index}｜{question['id']}｜{format_source(question)}", bold=True, color="1F4D78", size=9.5)
        set_cell_text(table.cell(1, 0), question_text(question), size=9.3)
        answer = f"参考答案：{format_answer(question)}"
        if question.get("explanation"):
            answer += f"\n解析：{question['explanation']}"
        set_cell_text(table.cell(2, 0), answer, size=9.3)


def add_exercise_index(doc: Document, questions: list[dict]) -> None:
    add_heading(doc, "原题训练索引", 2)
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    headers = ["题号", "来源", "题型/状态", "题干或要求"]
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, header, bold=True, color="1F4D78", size=8.8)
    for question in questions:
        row = table.add_row().cells
        status = "待核对答案" if question.get("answerPending") else "已配答案"
        if question.get("sourceTextMissing"):
            status += "；PDF文本缺失"
        values = [question["id"], format_source(question), f"{question['type']} / {status}", question_text(question)]
        for cell, value in zip(row, values):
            set_cell_text(cell, value, size=7.8)


def add_topic_guide(doc: Document, index: int, topic: dict, questions: list[dict]) -> None:
    guide = TOPIC_GUIDES[topic["id"]]
    add_heading(doc, f"第{index}讲  {guide['title']}", 1)
    add_label_box(doc, "目标要求", guide["targets"], fill="FFF3D8")
    source_counts = Counter(f"{q['source']['sourceId']} {q['source']['section']}" for q in questions)
    add_label_box(
        doc,
        "考情分析",
        [
            f"本讲纳入 PDF 原题 {len(questions)} 道。",
            "来源：" + "；".join(f"{key}={value}" for key, value in source_counts.most_common()),
        ],
        fill="EAF4EE",
    )
    image_path = make_network_image(topic["id"], guide)
    doc.add_picture(str(image_path), width=Inches(6.6))
    add_heading(doc, "分层 · 必备基础知识梳理", 2)
    add_bullets(doc, guide["knowledge"])
    add_label_box(doc, "特别提醒", guide["reminders"], fill="F4F6F9")
    add_examples(doc, questions)
    add_exercise_index(doc, questions)
    if index:
        doc.add_page_break()


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    questions = json.loads((DATA_DIR / "full_question_bank.json").read_text(encoding="utf-8"))
    topics = json.loads((DATA_DIR / "topics.json").read_text(encoding="utf-8"))
    topic_order = [topic for topic in topics if topic["id"] in TOPIC_GUIDES]
    by_topic: dict[str, list[dict]] = defaultdict(list)
    for question in questions:
        by_topic[question.get("topic", "uncategorized")].append(question)
    for topic_questions in by_topic.values():
        topic_questions.sort(key=lambda q: (q["source"]["sourceId"], q["source"]["section"], int(q["source"]["question"])))
    doc = Document()
    configure_document(doc)
    add_cover(doc, questions, topic_order)
    for index, topic in enumerate(topic_order, start=1):
        add_topic_guide(doc, index, topic, by_topic.get(topic["id"], []))
    output = OUT_DIR / "data_science_final_study_workbook.docx"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
